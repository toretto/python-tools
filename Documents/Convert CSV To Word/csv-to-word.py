from docx import Document
import csv

document = Document()
document.add_heading('BIO 2.0 Maatregelen')
document.add_paragraph('Dit document bevat de vertaling van de BIO 2.0 maatregelen naar de ISO 27001:2002 standaard')
document.add_page_break()

with open('data.csv', 'r') as file:
    with open('target.csv', 'w') as target:
        writer = csv.writer(target,delimiter=";")
        reader = csv.reader(file,delimiter=";")
        for row in reader:
            document.add_heading(f'Opmerking Control Bio 1.0.4zv {row[0]}',level=2)
            document.add_heading(f'ISO 27001:2022 control {row[3]}',level=3)
            document.add_heading(f'Control titel: {row[0]}',level=3)
            document.add_heading('Control',level=4)
            document.add_paragraph(f'{row[4]}')
            document.add_heading('Doel',level=4)
            document.add_paragraph(f'{row[5]}')
            document.add_heading('Overheidsmaatregel',level=4)
            document.add_paragraph(f'{row[10]}')
            document.add_heading('Verantwoordelijke',level=4)
            document.add_paragraph(f'{row[11]}')
            document.add_heading('Type control',level=4)
            document.add_paragraph(f'{row[12]}')
            document.add_heading('Informatiebeveiligingseigenschappen',level=4)
            document.add_paragraph(f'{row[13]}')
            document.add_heading('Cybersecurity Concepten',level=4)
            document.add_paragraph(f'{row[14]}')
            document.add_heading('Operationele Capaciteiten',level=4)
            document.add_paragraph(f'{row[15]}')
            document.add_heading('Beveiligingsdomeinen',level=4)
            document.add_paragraph(f'{row[16]}')
            document.add_heading('Opmerking Control Bio 1.04rv',level=4)
            document.add_paragraph(f'{row[6]}')
            document.add_page_break()

        document.save('BIO-20.docx')